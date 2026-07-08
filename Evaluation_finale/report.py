from docx import Document

document = Document()

document.add_heading("Rapport de conversation - CrediBot", level=1)

def add_conversation(question, answer):
    document.add_heading("Question", level=2)
    document.add_paragraph(question)

    document.add_heading("Réponse", level=2)
    document.add_paragraph(answer)

    document.add_page_break()

def save_report(filename="CrediBot_Report.docx"):
    document.save(filename)